import argparse
import os
import sys
import subprocess
import time
import re
import shutil
import pandas as pd
import markdown
import mammoth
from markdownify import markdownify as md_convert
from PIL import Image
from docx2pdf import convert as docx_to_pdf_tool
from docx import Document
from bs4 import BeautifulSoup
from google import genai
from dotenv import load_dotenv

load_dotenv()

# --- PROGRESS BAR HELPER ---
class ProgressBar:
    def __init__(self, total, desc="", width=80, unit="it"):
        self.total = total
        self.desc = desc
        self.width = width
        self.unit = unit
        self.start_time = time.time()
        
    def _format_time(self, seconds):
        if seconds is None or seconds < 0:
            return "--:--"
        m, s = divmod(int(seconds), 60)
        h, m = divmod(m, 60)
        if h > 0:
            return f"{h}:{m:02d}:{s:02d}"
        return f"{m:02d}:{s:02d}"

    def _format_value(self, val):
        if self.unit == "B":
            original_val = val
            for unit in ['', 'K', 'M', 'G', 'T']:
                if abs(val) < 1024.0:
                    if unit == '': return f"{val:.0f}"
                    if val >= 100: return f"{val:.0f}{unit}"
                    if val >= 10: return f"{val:.1f}{unit}"
                    return f"{val:.2f}{unit}"
                val /= 1024.0
            return f"{val:.1f}P"
        return f"{val:.2f}" if isinstance(val, float) else str(val)

    def update(self, current):
        """Updates the progress bar."""
        now = time.time()
        elapsed = now - self.start_time
        
        progress = current / self.total if self.total > 0 else 0
        progress = min(max(progress, 0), 1)
        percent = progress * 100
        
        rate = current / elapsed if elapsed > 0 else 0
        remaining = (self.total - current) / rate if rate > 0 else 0
        
        elapsed_str = self._format_time(elapsed)
        remaining_str = self._format_time(remaining)
        
        curr_str = self._format_value(current)
        total_str = self._format_value(self.total)
        
        if self.unit == "B":
            rate_str = f"{self._format_value(rate)}iB/s"
        else:
            rate_str = f"{rate:.2f}{self.unit}/s"
            
        # Format: 30%|███████████▎                          | 432M/1.42G [00:20<00:49, 21.7MiB/s]
        prefix = f"{int(percent)}%|"
        if self.desc:
            prefix = f"{self.desc}: {prefix}"
            
        suffix = f"| {curr_str}/{total_str} [{elapsed_str}<{remaining_str}, {rate_str}]"
        
        bar_width = self.width - len(prefix) - len(suffix)
        if bar_width < 10: bar_width = 10
        
        filled_len = bar_width * progress
        n_full = int(filled_len)
        remainder = filled_len - n_full
        
        fractions = [" ", "▏", "▎", "▍", "▌", "▋", "▊", "▉"]
        
        bar = "█" * n_full
        if n_full < bar_width:
            idx = int(remainder * 8)
            bar += fractions[idx]
            bar += " " * (bar_width - n_full - 1)
        
        line = f"{prefix}{bar}{suffix}"
        sys.stdout.write(f"\r\033[K{line}")
        sys.stdout.flush()

    def finish(self):
        sys.stdout.write("\n")
        sys.stdout.flush()

    def finish(self):
        sys.stdout.write("\n")
        sys.stdout.flush()

def safe_print(msg):
    """Prints a message safely, clearing the line first."""
    sys.stdout.write(f"\r\033[K{msg}\n")
    sys.stdout.flush()

# --- FFMPEG HELPER ---
def get_duration(input_path):
    """Returns duration in seconds using ffmpeg."""
    try:
        cmd = ["ffmpeg", "-i", input_path]
        result = subprocess.run(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE, text=True)
        # Search for Duration: HH:MM:SS.mm
        match = re.search(r"Duration: (\d{2}):(\d{2}):(\d{2}\.\d+)", result.stderr)
        if match:
            h, m, s = map(float, match.groups())
            return h * 3600 + m * 60 + s
    except Exception:
        pass
    return 0

def run_ffmpeg_with_progress(cmd, desc):
    """Runs ffmpeg command and updates progress bar."""
    # 1. Get Duration first from input (assumes input is after -i)
    input_file = None
    try:
        idx = cmd.index("-i")
        if idx + 1 < len(cmd):
            input_file = cmd[idx + 1]
    except ValueError:
        pass
        
    duration = get_duration(input_file) if input_file else 0
    
    if duration <= 0:
        # No duration known, just run silently or with generic wait
        safe_print(f"   (No duration info, running silently...)")
        subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.PIPE)
        return

    # 2. Run with -progress pipe:1
    # We append flags to ensure progress is output to stdout
    full_cmd = cmd + ["-progress", "pipe:1", "-nostats"]
    
    process = subprocess.Popen(
        full_cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        bufsize=1
    )
    
    pb = ProgressBar(total=duration, desc=desc, unit="s")
    
    while True:
        line = process.stdout.readline()
        if not line and process.poll() is not None:
            break
        
        if line:
            # Parse out_time_us=1234567
            if "out_time_us=" in line:
                try:
                    parts = line.strip().split("=")
                    if len(parts) == 2:
                        us = int(parts[1])
                        current_sec = us / 1_000_000.0
                        pb.update(current_sec)
                except ValueError:
                    pass
    
    pb.finish()
    
    if process.returncode != 0:
        err_out = process.stderr.read()
        safe_print(f"❌ FFmpeg Error: {err_out[-200:] if err_out else 'Unknown error'}")

# --- AI ENHANCEMENT ---
def apply_gemini_enhancement(text_content):
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        safe_print("❌ Error: GEMINI_API_KEY not found.")
        return text_content 

    safe_print("   ✨ Sending to Gemini...")
    try:
        client = genai.Client(api_key=api_key)
        
        prompt = (
            "You are a professional editor. Enhance the text below. "
            "Fix grammar, improve tone, and format clearly. Return ONLY the text.\n\n"
            f"{text_content}"
        )
        
        response = client.models.generate_content(
            model='gemini-2.0-flash', 
            contents=prompt
        )
        return response.text
    except Exception as e:
        safe_print(f"   ⚠️ Gemini Error: {e}")
        return text_content

# --- 1. IMAGE LOGIC ---
def convert_image(input_path, output_path, reduce_mode=False, enhance_mode=False):
    try:
        # Dummy progress for image since PIL blocks
        pb = ProgressBar(total=1, desc="Image", unit="img")
        pb.update(0)
        
        with Image.open(input_path) as img:
            # FIX: Force RGB for JPEGs
            if output_path.lower().endswith(('.jpg', '.jpeg')):
                if img.mode in ('RGBA', 'LA', 'P'): 
                    img = img.convert('RGB')
            
            if reduce_mode:
                img.save(output_path, optimize=True, quality=60)
            else:
                img.save(output_path)
        
        pb.update(1)
        pb.finish()
    except Exception as e:
        safe_print(f"❌ Image Error: {e}")

# --- 2. VIDEO LOGIC ---
def convert_video(input_path, output_path, reduce_mode=False, enhance_mode=False):
    try:
        if subprocess.call(["which", "ffmpeg"], stdout=subprocess.DEVNULL) != 0:
            safe_print("❌ Error: FFmpeg is not installed.")
            return
        
        cmd = ["ffmpeg", "-i", input_path, "-y"]
        
        if reduce_mode:
            cmd.extend(["-vcodec", "libx264", "-crf", "28", "-preset", "fast"])
        
        # MKV Logic: Preserve all streams (video, audio, subtitles)
        if output_path.lower().endswith('.mkv'):
            cmd.extend(["-map", "0"])
            # If not reducing, default to copying streams for speed/quality if codecs allow?
            # For now, just ensure we map everything.

        cmd.append(output_path)
        
        run_ffmpeg_with_progress(cmd, desc="Video")
        
    except Exception as e:
        safe_print(f"❌ Video Error: {e}")

# --- 3. AUDIO LOGIC ---
def convert_audio(input_path, output_path, reduce_mode=False, enhance_mode=False):
    try:
        if subprocess.call(["which", "ffmpeg"], stdout=subprocess.DEVNULL) != 0:
            safe_print("❌ Error: FFmpeg is not installed.")
            return

        cmd = ["ffmpeg", "-i", input_path, "-y"]
        out_ext = os.path.splitext(output_path)[1].lower()

        if out_ext == '.mp3':
            cmd.extend(["-c:a", "libmp3lame"])
        elif out_ext == '.ogg':
            cmd.extend(["-c:a", "libvorbis"])

        if reduce_mode:
            if out_ext == '.flac':
                cmd.extend(["-compression_level", "12"])
            elif out_ext == '.mp3':
                cmd.extend(["-b:a", "128k"])
            elif out_ext in ['.m4a', '.aac']:
                cmd.extend(["-c:a", "aac", "-b:a", "128k"])
            elif out_ext == '.ogg':
                cmd.extend(["-q:a", "3"])
        
        cmd.append(output_path)
        
        run_ffmpeg_with_progress(cmd, desc="Audio")
        
    except Exception as e:
        safe_print(f"❌ Audio Error: {e}")

# --- 4. DOCS LOGIC ---
def convert_docs(input_path, output_path, reduce_mode=False, enhance_mode=False):
    try:
        in_ext = os.path.splitext(input_path)[1].lower()
        out_ext = os.path.splitext(output_path)[1].lower()
        content = ""

        safe_print(f"   Reading {in_ext}...")

        # Extract
        if in_ext in ['.md', '.txt', '.mdx']:
            with open(input_path, 'r', encoding='utf-8') as f: content = f.read()
        elif in_ext == '.docx':
            with open(input_path, "rb") as docx_file:
                content = mammoth.convert_to_html(docx_file).value
                if out_ext in ['.md', '.mdx']: content = md_convert(content)

        # Enhance
        if enhance_mode and content:
            content = apply_gemini_enhancement(content)

        safe_print(f"   Writing {out_ext}...")

        # Write
        if out_ext == '.html':
            html_content = markdown.markdown(content) if in_ext in ['.md', '.mdx'] else content
            with open(output_path, 'w', encoding='utf-8') as f: f.write(html_content)
        elif out_ext in ['.md', '.txt', '.mdx']:
            with open(output_path, 'w', encoding='utf-8') as f: f.write(content)
        elif out_ext == '.pdf' and in_ext == '.docx':
            docx_to_pdf_tool(os.path.abspath(input_path), os.path.abspath(output_path))
        elif out_ext == '.docx':
            doc = Document()
            html_content = markdown.markdown(content)
            soup = BeautifulSoup(html_content, 'html.parser')

            def process_inline(paragraph, element):
                if element.name is None:
                    paragraph.add_run(element.string)
                else:
                    if element.name == 'strong' or element.name == 'b':
                        run = paragraph.add_run(element.get_text())
                        run.bold = True
                    elif element.name == 'em' or element.name == 'i':
                        run = paragraph.add_run(element.get_text())
                        run.italic = True
                    elif element.name == 'code':
                        run = paragraph.add_run(element.get_text())
                        run.font.name = 'Courier New'
                    else:
                        paragraph.add_run(element.get_text())

            def process_block(element):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(element.name[1])
                    p = doc.add_heading(level=level)
                    for child in element.children:
                        process_inline(p, child)
                elif element.name == 'p':
                    p = doc.add_paragraph()
                    for child in element.children:
                        process_inline(p, child)
                elif element.name in ['ul', 'ol']:
                    is_ordered = element.name == 'ol'
                    style = 'List Number' if is_ordered else 'List Bullet'
                    for li in element.find_all('li', recursive=False):
                        p = doc.add_paragraph(style=style)
                        for child in li.children:
                            process_inline(p, child)

            for element in soup.body.children if soup.body else soup.children:
                if element.name:
                    process_block(element)

            doc.save(output_path)

    except Exception as e:
        safe_print(f"❌ Doc Error: {e}")

# --- 5. DATA LOGIC ---
def convert_data(input_path, output_path, reduce_mode=False, enhance_mode=False):
    try:
        safe_print("   Processing data...")
        in_ext = os.path.splitext(input_path)[1].lower()
        if in_ext == '.csv': df = pd.read_csv(input_path)
        elif in_ext == '.json': df = pd.read_json(input_path)
        elif in_ext in ['.xlsx', '.xls']: df = pd.read_excel(input_path)
        else: return

        out_ext = os.path.splitext(output_path)[1].lower()
        if out_ext == '.csv': df.to_csv(output_path, index=False)
        elif out_ext == '.json': df.to_json(output_path, orient='records', indent=4)
        elif out_ext == '.xlsx': df.to_excel(output_path, index=False)
    except Exception as e:
        safe_print(f"❌ Data Error: {e}")

# --- DISPATCHER ---
def get_converter(input_ext, output_ext):
    img_exts = {'.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff'}
    vid_exts = {'.mp4', '.mkv', '.mov', '.avi', '.webm', '.ogv', '.ogm'}
    audio_exts = {'.mp3', '.wav', '.aac', '.flac', '.ogg', '.m4a', '.opus'}
    data_exts = {'.csv', '.json', '.xlsx'}
    doc_exts  = {'.docx', '.pdf', '.md', '.html', '.txt', '.mdx'}

    if input_ext in img_exts and output_ext in img_exts: return convert_image
    if input_ext in vid_exts and output_ext in vid_exts: return convert_video
    if input_ext in vid_exts and output_ext in audio_exts: return convert_audio
    if input_ext in audio_exts and output_ext in audio_exts: return convert_audio
    if input_ext in data_exts and output_ext in data_exts: return convert_data
    if input_ext in doc_exts and output_ext in doc_exts: return convert_docs
    return None

# --- CLI ENTRY ---
def main():
    parser = argparse.ArgumentParser(description="Universal Converter")
    parser.add_argument("args", nargs="+", help="Files + Output Path")
    parser.add_argument("-t", "--to", help="Target format")
    parser.add_argument("-r", "--reduce", action="store_true", help="Compress file size")
    parser.add_argument("-e", "--enhance", action="store_true", help="AI Enhancement")
    
    args = parser.parse_args()
    
    inputs = args.args[:-1]
    destination = args.args[-1]

    if len(args.args) == 1 and args.to:
        inputs = [args.args[0]]
        destination = os.path.dirname(args.args[0]) or "."
    elif len(inputs) == 0:
        print("Usage: converter <files> <output_folder> [options]")
        return

    output_dir = destination
    if not os.path.exists(output_dir) and len(inputs) > 0 and not os.path.splitext(destination)[1]:
        os.makedirs(output_dir)

    target_ext = f".{args.to.lstrip('.')}" if args.to else None
    
    total_files = len(inputs)
    print(f"🚀 Starting conversion of {total_files} files...")
    
    for i, input_file in enumerate(inputs, 1):
        if not os.path.exists(input_file):
            safe_print(f"⚠️  Missing: {input_file}")
            continue
        
        in_ext = os.path.splitext(input_file)[1].lower()
        final_target = target_ext if target_ext else os.path.splitext(destination)[1].lower()
        
        if os.path.isdir(destination):
            base = os.path.splitext(os.path.basename(input_file))[0]
            out_file = os.path.join(destination, base + final_target)
        else:
            out_file = destination

        print(f"\n[{i}/{total_files}] File: {os.path.basename(input_file)}")

        func = get_converter(in_ext, final_target)
        if func:
            func(input_file, out_file, reduce_mode=args.reduce, enhance_mode=args.enhance)
        else:
            safe_print(f"⚠️  No logic for {in_ext} -> {final_target}")

    print("\n✅ All tasks completed.")

if __name__ == "__main__":
    main()